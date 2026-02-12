import os
import urllib.parse
from docx import Document
from docx.shared import Pt, RGBColor
from docx.opc.constants import RELATIONSHIP_TYPE as RELATIONSHIP_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# --- CONFIGURAZIONE ---
OUTPUT_DIR = "outputs"

def _ensure_output_dir():
    """
    Crea la cartella outputs se non esiste.
    Usa il percorso assoluto per evitare errori se lanciato da altre posizioni.
    """
    base_path = os.getcwd()
    full_output_path = os.path.join(base_path, OUTPUT_DIR)
    
    if not os.path.exists(full_output_path):
        os.makedirs(full_output_path)
    return full_output_path

def generate_gmaps_search_link(name, address):
    """Genera un link di ricerca universale (evita errori 404)."""
    if not name and not address:
        return None
    
    query = f"{name} {address}".strip()
    # Codifica URL sicura
    safe_query = urllib.parse.quote(query)
    # Usa l'API di ricerca universale ufficiale
    return f"https://www.google.com/maps/search/?api=1&query={safe_query}"

def _format_terminal_link(label, url):
    if not url:
        return label
    # Plain URL for maximum compatibility across terminals
    return f"{url}"

def _add_docx_hyperlink(paragraph, text, url):
    """
    Add a clickable hyperlink to a docx paragraph.
    """
    part = paragraph.part
    r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0000FF")
    rpr.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.append(underline)

    run.append(rpr)
    text_elem = OxmlElement("w:t")
    text_elem.text = text
    run.append(text_elem)

    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    return hyperlink


def _top_flight_options(state, max_items=3):
    options = state.get("flight_options") or []
    if not isinstance(options, list):
        return []
    return options[:max_items]


def _selected_flight(state):
    options = _top_flight_options(state, max_items=1)
    if not options:
        return None
    return options[0]


def _format_flight_price(option):
    value = option.get("price_value")
    if value is None:
        return "n/d"
    try:
        return f"{float(value):.2f}"
    except (TypeError, ValueError):
        return "n/d"

def print_terminal_report(state):
    if not state.get('itinerary') and not state.get('is_approved'):
        print("\n" + "!"*60)
        print(" [!] EMERGENCY ABORT: VALIDATION FAILED")
        print(f" CAUSA: {state.get('critic_feedback')}")
        print("!"*60 + "\n")
        return
    
    print("\n" + "="*60)
    print(f" >>> ITINERARIO FINALE: {state.get('destination', 'Viaggio').upper()}")
    print("="*60)

    flight_summary = state.get("flight_summary")
    selected = _selected_flight(state)
    if flight_summary or selected:
        print("\n [FLIGHT]")
        if flight_summary:
            print(f"   {flight_summary}")
        if selected:
            title = selected.get("title", "N/D")
            url = selected.get("url", "")
            price = _format_flight_price(selected)
            dep_date = selected.get("depart_date", "n/d")
            dep_time = selected.get("depart_time", "n/d")
            print(f"   1. {title} | prezzo stimato: {price}")
            print(f"      data: {dep_date} | orario: {dep_time}")
            if url:
                print(f"      {url}")
            ret_title = selected.get("return_title")
            if ret_title and ret_title != "n/d":
                ret_price = selected.get("return_price_value")
                try:
                    ret_price_text = f"{float(ret_price):.2f}" if ret_price is not None else "n/d"
                except (TypeError, ValueError):
                    ret_price_text = "n/d"
                ret_date = selected.get("return_depart_date", "n/d")
                ret_time = selected.get("return_depart_time", "n/d")
                ret_url = selected.get("return_url", "")
                print(f"   2. {ret_title} | prezzo stimato: {ret_price_text}")
                print(f"      data: {ret_date} | orario: {ret_time}")
                if ret_url:
                    print(f"      {ret_url}")
    
    itinerary = state.get('itinerary', [])
    for day in itinerary:
        print(f"\n [+] Giorno {day['day_number']}: {day['focus']}")
        print("-" * 40)
        for p in day.get('places', []):
            name = p.get('name', 'Senza nome')
            rating = p.get('rating', 'N/A')
            address = p.get('address', '')
            
            print(f"   [o] {name} (Score: {rating})") 
            if address:
                print(f"       Addr: {address}")
            
            link = generate_gmaps_search_link(name, address)
            if link:
                print("       Link:")
                print(f"       {_format_terminal_link('Apri su Maps', link)}")
            print("       " + "." * 20)

def generate_html_report(state):
    """Genera un report HTML visivamente ricco."""
    output_dir = _ensure_output_dir()
    
    dest = state.get('destination', 'Viaggio')
    filename = f"viaggio_{dest.replace(' ', '_').lower()}.html"
    filepath = os.path.join(output_dir, filename)
    
    html = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <title>Viaggio a {dest}</title>
        <meta charset="utf-8">
        <style>
            body {{ font-family: 'Segoe UI', sans-serif; padding: 20px; background: #f0f2f5; color: #333; }}
            h1 {{ color: #2c3e50; text-align: center; border-bottom: 2px solid #3498db; padding-bottom: 10px; }}
            .day-card {{ background: white; padding: 20px; margin-bottom: 20px; border-radius: 10px; box-shadow: 0 4px 6px rgba(0,0,0,0.1); }}
            .flight-card {{ background: #eef6ff; padding: 16px; margin-bottom: 20px; border-radius: 10px; border-left: 6px solid #3498db; }}
            h2 {{ color: #e67e22; }}
            .place {{ margin-top: 15px; padding: 10px; background: #f9f9f9; border-left: 5px solid #3498db; border-radius: 4px; }}
            .place-name {{ font-size: 1.1em; font-weight: bold; }}
            .rating {{ color: #f1c40f; font-weight: bold; }}
            .address {{ font-style: italic; color: #555; margin: 5px 0; }}
            a.btn {{ display: inline-block; margin-top: 5px; background: #3498db; color: white; padding: 5px 10px; text-decoration: none; border-radius: 4px; font-size: 0.9em; }}
            a.btn:hover {{ background: #2980b9; }}
        </style>
    </head>
    <body>
        <h1>✈️ Itinerario: {dest.upper()}</h1>
    """

    flight_summary = state.get("flight_summary")
    selected = _selected_flight(state)
    if flight_summary or selected:
        html += "<div class='flight-card'><h2>Volo suggerito</h2>"
        if flight_summary:
            html += f"<p>{flight_summary}</p>"
        if selected:
            html += "<ul>"
            title = selected.get("title", "N/D")
            url = selected.get("url", "")
            price = _format_flight_price(selected)
            dep_date = selected.get("depart_date", "n/d")
            dep_time = selected.get("depart_time", "n/d")
            if url:
                html += (
                    f"<li>{title} | prezzo stimato: {price} | data: {dep_date} | orario: {dep_time} "
                    f"- <a href='{url}' target='_blank'>Link</a></li>"
                )
            else:
                html += f"<li>{title} | prezzo stimato: {price} | data: {dep_date} | orario: {dep_time}</li>"

            ret_title = selected.get("return_title")
            if ret_title and ret_title != "n/d":
                ret_price = selected.get("return_price_value")
                try:
                    ret_price_text = f"{float(ret_price):.2f}" if ret_price is not None else "n/d"
                except (TypeError, ValueError):
                    ret_price_text = "n/d"
                ret_date = selected.get("return_depart_date", "n/d")
                ret_time = selected.get("return_depart_time", "n/d")
                ret_url = selected.get("return_url", "")
                if ret_url:
                    html += (
                        f"<li>{ret_title} | prezzo stimato: {ret_price_text} | data: {ret_date} | orario: {ret_time} "
                        f"- <a href='{ret_url}' target='_blank'>Link</a></li>"
                    )
                else:
                    html += (
                        f"<li>{ret_title} | prezzo stimato: {ret_price_text} | data: {ret_date} | orario: {ret_time}</li>"
                    )
            html += "</ul>"
        html += "</div>"
    
    for day in state.get('itinerary', []):
        html += f"<div class='day-card'><h2>Giorno {day['day_number']}: {day['focus']}</h2>"
        for p in day.get('places', []):
            name = p.get('name', 'Senza nome')
            address = p.get('address', '')
            rating = p.get('rating', 'N/A')
            link = generate_gmaps_search_link(name, address)
            
            html += f"""
            <div class='place'>
                <div class='place-name'>{name} <span class='rating'>★ {rating}</span></div>
                <div class='address'>{address}</div>
            """
            if link:
                html += f"<a href='{link}' class='btn' target='_blank'>Vedi su Maps</a>"
            html += "</div>"
        html += "</div>"
        
    html += "</body></html>"
    
    with open(filepath, "w", encoding="utf-8") as f:
        f.write(html)
        
    return filepath

def generate_docx_report(state):
    """Genera un file Word ben formattato."""
    output_dir = _ensure_output_dir()
    
    destination = state.get('destination', 'Viaggio')
    filename = f"viaggio_{destination.replace(' ', '_').lower()}.docx"
    filepath = os.path.join(output_dir, filename)
    
    doc = Document()
    
    # Titolo Principale
    title = doc.add_heading(f'Itinerario: {destination.upper()}', 0)
    title.alignment = 1 # Center

    doc.add_paragraph(f"Ecco il tuo piano di viaggio generato dall'AI per {destination}.")

    flight_summary = state.get("flight_summary")
    selected = _selected_flight(state)
    if flight_summary or selected:
        doc.add_heading("Volo suggerito", level=1)
        if flight_summary:
            doc.add_paragraph(flight_summary)
        if selected:
            title_text = selected.get("title", "N/D")
            price = _format_flight_price(selected)
            url = selected.get("url", "")
            dep_date = selected.get("depart_date", "n/d")
            dep_time = selected.get("depart_time", "n/d")
            doc.add_paragraph(
                f"1. {title_text} | prezzo stimato: {price} | data: {dep_date} | orario: {dep_time}"
            )
            if url:
                p_link = doc.add_paragraph(style='List Bullet')
                _add_docx_hyperlink(p_link, "Apri offerta volo", url)

            ret_title = selected.get("return_title")
            if ret_title and ret_title != "n/d":
                ret_price = selected.get("return_price_value")
                try:
                    ret_price_text = f"{float(ret_price):.2f}" if ret_price is not None else "n/d"
                except (TypeError, ValueError):
                    ret_price_text = "n/d"
                ret_date = selected.get("return_depart_date", "n/d")
                ret_time = selected.get("return_depart_time", "n/d")
                ret_url = selected.get("return_url", "")
                doc.add_paragraph(
                    f"2. {ret_title} | prezzo stimato: {ret_price_text} | data: {ret_date} | orario: {ret_time}"
                )
                if ret_url:
                    p_link = doc.add_paragraph(style='List Bullet')
                    _add_docx_hyperlink(p_link, "Apri offerta volo ritorno", ret_url)

    for day in state.get('itinerary', []):
        # Intestazione Giorno
        doc.add_heading(f"Giorno {day['day_number']}: {day['focus']}", level=1)
        
        for p in day.get('places', []):
            name = p.get('name', 'Senza nome')
            rating = p.get('rating', 'N/A')
            address = p.get('address', '')
            
            # Nome Luogo 
            p_para = doc.add_paragraph()
            runner = p_para.add_run(f"{name}")
            runner.bold = True
            runner.font.size = Pt(12)
            runner.font.color.rgb = RGBColor(0, 51, 102) # Blu scuro
            
            # Dettagli
            doc.add_paragraph(f"   ⭐ Rating: {rating}")
            doc.add_paragraph(f"   Indirizzo: {address}")
            
            # Link Maps
            link = generate_gmaps_search_link(name, address)
            if link:
                
                p_link = doc.add_paragraph(style='List Bullet')
                _add_docx_hyperlink(p_link, "Apri su Maps", link)
            
            doc.add_paragraph("_" * 40) # Separatore

    doc.save(filepath)
    return filepath
